import requests
from bs4 import BeautifulSoup
import datetime
import pandas as pd
import argparse
import os
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


def fetch_loan_rates():
    """爬取中国银行贷款利率数据"""
    url = "https://www.bankofchina.com/fimarkets/lilv/fd32/201310/t20131031_2591219.html"

    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(url, headers=headers)
        response.encoding = 'utf-8'  # 确保中文正确显示

        if response.status_code == 200:
            soup = BeautifulSoup(response.text, 'html.parser')

            # 首先查找页面中的所有表格
            tables = soup.find_all('table')

            # 查找包含"LPR"的表格
            lpr_table = None
            for table in tables:
                if 'LPR' in table.get_text():
                    lpr_table = table
                    break

            # 如果找到了表格，解析其中的数据
            if lpr_table:
                rows = lpr_table.find_all('tr')
                all_text = []
                for row in rows:
                    cells = row.find_all(['td', 'th'])
                    if cells:
                        row_text = ' '.join(cell.get_text(strip=True)
                                            for cell in cells)
                        all_text.append(row_text)

            with open('LPR_Data.txt', 'w', encoding='utf-8') as f:
                f.write('\n'.join(all_text))
                f.close()

            return True
        else:
            print(f"请求失败，状态码: {response.status_code}")
            return False

    except Exception as e:
        print(f"发生错误: {e}")
        return False


def load_lpr_data(file_path='LPR_Data.txt', check_update=True):
    """从文件加载LPR数据，并转换为DataFrame，如果数据过旧则自动更新"""
    try:
        # 尝试加载现有数据
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        # 跳过标题行
        data = []
        for line in lines[1:]:  # 跳过第一行（标题行）
            parts = line.strip().split()
            if len(parts) >= 3:
                date = parts[0]
                one_year_rate = float(parts[1].replace('%', '')) / 100
                five_year_rate = float(parts[2].replace('%', '')) / 100
                data.append([date, one_year_rate, five_year_rate])

        # 创建DataFrame
        df = pd.DataFrame(
            data, columns=['date', 'one_year_rate', 'five_year_rate'])
        df['date'] = pd.to_datetime(df['date'])
        df = df.sort_values('date')  # 确保按日期排序

        # 只在第一次调用时检查更新，避免死循环
        if check_update and len(df) > 0:
            today = datetime.datetime.now()
            last_date = df['date'].iloc[-1]

            if today > last_date + datetime.timedelta(days=30):
                print("数据已超过30天，正在更新...")
                # 调用更新函数
                update_successful = fetch_loan_rates()

                # 只有在更新成功的情况下才重新加载
                if update_successful:
                    # 设置check_update为False，避免死循环
                    return load_lpr_data(file_path, check_update=False)
                else:
                    print("更新失败，使用当前数据")

        return df
    except Exception as e:
        print(f"加载LPR数据时发生错误: {e}")
        return None


def calculate_days(start_date, end_date, gap):
    """根据给定的开始日期、结束日期和计算方式计算天数
    
    Args:
        start_date: 开始日期
        end_date: 结束日期
        gap: 天数计算方式，'both'表示两头都算，'no_tail'表示算头不算尾
    
    Returns:
        int: 计算得出的天数
    """
    if gap == "both":
        # 两头都算：(end_date - start_date).days + 1
        return (end_date - start_date).days + 1
    else:  # no_tail
        # 算头不算尾：(end_date - start_date).days
        return (end_date - start_date).days


def print_interest_details(amount, start_date, end_date, lpr_data, term='one_year', mag=1, gap="no_tail", day_count=365):
    """打印合并后的利息详情，将相同利率的时间段合并计算"""
    rate_column = 'one_year_rate' if term == 'one_year' else 'five_year_rate'
    mag_str = f" × {mag}" if mag > 1 else ""

    # 检查开始日期是否在LPR数据范围内
    if start_date < lpr_data['date'].min():
        print(
            f"警告: 开始日期早于LPR数据范围（最早数据日期：{lpr_data['date'].min().strftime('%Y-%m-%d')}）")
        return None

    # 如果结束日期超出LPR数据范围，记录一个标志，后面会使用最新的LPR利率
    end_date_exceeds = False
    if end_date > lpr_data['date'].max():
        print(
            f"注意: 结束日期晚于LPR数据范围（最新数据日期：{lpr_data['date'].max().strftime('%Y-%m-%d')}），将使用最新的LPR利率计算剩余期间")
        end_date_exceeds = True

    # 获取开始日期之前的最后一个LPR变更
    initial_rate_row = lpr_data[lpr_data['date'] <= start_date].iloc[-1]
    initial_rate = initial_rate_row[rate_column]
    initial_date = initial_rate_row['date']

    # 获取计算期间内的所有LPR变更日期
    rate_changes = lpr_data[(lpr_data['date'] > start_date)
                            & (lpr_data['date'] <= end_date)]

    # 初始化利息总额和存储计算结果
    total_interest = 0
    calculation_results = []

    # 打印详情标题
    print(f"\n{'=' * 60}")
    print(f"计算本金金额: {amount:,.2f} 元")
    print(
        f"计算期间: {start_date.strftime('%Y-%m-%d')} 至 {end_date.strftime('%Y-%m-%d')}")
    print(f"LPR期限: {'一年期' if term == 'one_year' else '五年期以上'}")
    print(f"约定倍数: {mag} 倍")
    print(f"天数算法: {'两头都算' if gap == 'both' else '算头不算尾'}")
    print(f"计息基础: 每年{day_count}天")
    print(f"{'=' * 60}")
    print(f"{'开始日期':<12}{'结束日期':<12}{'天数':<6}{'适用LPR':<10}{'计算金额':<15}")
    print(f"{'-' * 60}")

    # 处理计算逻辑
    current_date = start_date  # 初始日期
    current_rate = initial_rate  # 初始利率

    # 为合并相同利率的计算段做准备
    rate_segments = []
    
    # 1. 先收集所有利率变更点
    for idx, row in rate_changes.iterrows():
        rate_segments.append({
            'start_date': current_date,
            'end_date': row['date'] - datetime.timedelta(days=1),  # 利率变更前一天
            'rate': current_rate
        })
        
        # 更新当前日期和利率
        current_date = row['date']
        current_rate = row[rate_column]
    
    # 添加最后一个阶段
    rate_segments.append({
        'start_date': current_date,
        'end_date': end_date,
        'rate': current_rate
    })
    
    # 2. 合并相同利率的段
    merged_segments = []
    current_segment = None
    
    for segment in rate_segments:
        if current_segment is None:
            current_segment = segment.copy()
        elif current_segment['rate'] == segment['rate']:
            # 如果利率相同，合并日期范围
            current_segment['end_date'] = segment['end_date']
        else:
            # 利率不同，添加当前段并开始新的段
            merged_segments.append(current_segment)
            current_segment = segment.copy()
    
    # 添加最后一个段
    if current_segment is not None:
        merged_segments.append(current_segment)
    
    # 3. 计算每个合并段的利息
    for i, segment in enumerate(merged_segments):
        # Check if this is the last segment
        is_last_segment = (i == len(merged_segments) - 1)
        
        days = calculate_days(segment['start_date'], segment['end_date'], gap)
        # Add 1 to days count for all segments except the last one
        if not is_last_segment:
            days += 1
            
        interest = amount * segment['rate'] * mag * days / day_count
        total_interest += interest
        
        print(f"{segment['start_date'].strftime('%Y-%m-%d'):<12}{segment['end_date'].strftime('%Y-%m-%d'):<12}{days:<6}{segment['rate'] * 100:.2f}%{mag_str}{interest:>15,.2f}")
        
        calculation_results.append({
            'start_date': segment['start_date'],
            'end_date': segment['end_date'],
            'days': days,
            'rate': segment['rate'],
            'note': mag_str,
            'interest': interest
        })
        
    # 计算并打印总天数
    total_days = sum(calc['days'] for calc in calculation_results)
    print(f"{'-' * 60}")
    print(f"{'总天数':<30}{total_days:>30}")
    print(f"{'总金额':<30}{total_interest:>30,.2f}")
    print(f"{'=' * 60}")

    # 返回计算结果和总利息
    return {
        'amount': amount,
        'start_date': start_date,
        'end_date': end_date,
        'term': term,
        'mag': mag,
        'gap': gap,
        'day_count': day_count,
        'calculations': calculation_results,
        'total_interest': total_interest
    }


def export_to_word(results, output_file='LPR计算报告.docx'):
    """将计算结果导出为Word文档"""
    doc = Document()

    # 全局默认字体
    doc.styles['Normal'].font.name = '仿宋'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    doc.styles['Normal'].font.size = Pt(12)
    doc.styles['Normal'].paragraph_format.line_spacing = 1.1
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)
    doc.styles['Normal'].paragraph_format.space_after = Pt(0)

    # 标题字体
    # 这一步很重要，清除样式，自己设置的字体名称才会生效
    doc.styles['Title']._element.rPr.rFonts.clear()
    doc.styles['Title'].font.name = '黑体'
    doc.styles['Title']._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    doc.styles['Title'].font.size = Pt(26)
    doc.styles['Title'].paragraph_format.space_before = Pt(0)
    doc.styles['Title'].paragraph_format.space_after = Pt(0)
    doc.styles['Title'].paragraph_format.line_spacing = 1.1

    doc.styles['Heading 1']._element.rPr.rFonts.clear()
    doc.styles['Heading 1'].font.name = '黑体'
    doc.styles['Heading 1']._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    doc.styles['Heading 1'].font.size = Pt(14)
    doc.styles['Heading 1'].paragraph_format.line_spacing = 1.1
    doc.styles['Heading 1'].paragraph_format.space_before = Pt(0)
    doc.styles['Heading 1'].paragraph_format.space_after = Pt(0)

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # 添加标题
    title = doc.add_heading('LPR计算报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加报告生成日期
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_paragraph.add_run(
        f"报告生成日期: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # 添加参数信息
    doc.add_heading('计算参数', level=1)

    params_table = doc.add_table(rows=7, cols=2)
    params_table.style = 'Table Grid'

    # 设置表格内容
    cells = params_table.rows[0].cells
    cells[0].text = '计算金额'
    cells[1].text = f"{results['amount']:,.2f} 元"

    cells = params_table.rows[1].cells
    cells[0].text = '计算期间'
    cells[1].text = f"{results['start_date'].strftime('%Y-%m-%d')} 至 {results['end_date'].strftime('%Y-%m-%d')}"

    cells = params_table.rows[2].cells
    cells[0].text = 'LPR期限'
    cells[1].text = '一年期' if results['term'] == 'one_year' else '五年期以上'

    cells = params_table.rows[3].cells
    cells[0].text = '约定倍数'
    cells[1].text = f"{results['mag']} 倍"

    cells = params_table.rows[4].cells
    cells[0].text = '天数算法'
    cells[1].text = '两头都算' if results['gap'] == 'both' else '算头不算尾'

    cells = params_table.rows[5].cells
    cells[0].text = '计息基础'
    cells[1].text = f"每年{results['day_count']}天"
    # 计算总天数
    total_days = sum(calc['days'] for calc in results['calculations'])

    cells = params_table.rows[6].cells
    cells[0].text = '总天数'
    cells[1].text = f"{total_days} 天"

    # 增加一行显示总金额
    params_table.add_row()
    cells = params_table.rows[7].cells
    cells[0].text = '总金额'
    cells[1].text = f"{results['total_interest']:,.2f} 元"

    # 添加间隔
    doc.add_paragraph()

    # 添加计算明细
    doc.add_heading('分段计算明细', level=1)

    # 创建表格
    details_table = doc.add_table(
        rows=len(results['calculations']) + 2, cols=5)
    details_table.style = 'Table Grid'

    # 设置表头
    header_cells = details_table.rows[0].cells
    header_cells[0].text = '开始日期'
    header_cells[1].text = '结束日期'
    header_cells[2].text = '天数'
    header_cells[3].text = '适用LPR'
    header_cells[4].text = '利息金额'

    # 添加数据行
    for i, calc in enumerate(results['calculations']):
        row = details_table.rows[i + 1].cells
        row[0].text = calc['start_date'].strftime('%Y-%m-%d')
        row[1].text = calc['end_date'] .strftime('%Y-%m-%d')
        row[2].text = str(calc['days'])
        row[3].text = f"{calc['rate']:.2%}{calc['note']}"
        row[4].text = f"{calc['interest']:,.2f}"

    # 添加总计行
    total_row = details_table.rows[-1].cells
    total_row[0].text = '总计'
    total_row[0].merge(total_row[3])
    total_row[4].text = f"{results['total_interest']:,.2f}"

    # 添加间隔
    doc.add_paragraph()

    # 添加说明
    doc.add_heading('说明', level=1)
    doc.add_paragraph('1. 本报告基于中国人民银行公布的LPR（贷款市场报价利率）数据计算。')
    doc.add_paragraph('2. 计算采用实际天数和指定的年度天数（360或365）计算。')
    doc.add_paragraph('3. 当LPR变更时，利息计算会按照不同的时间段分别计算，相同LPR利率的时间段已合并显示。')
    doc.add_paragraph('4. 如果计算日期超出已公布的LPR数据范围，则使用最新的LPR利率。')

    # 保存文档
    doc.save(output_file)
    print(f"\n已将计算结果导出到Word文档: {os.path.abspath(output_file)}")


def parse_arguments():
    """解析命令行参数"""
    parser = argparse.ArgumentParser(description='根据LPR计算利息')
    parser.add_argument('--amount', type=float, required=True, help='贷款金额')
    parser.add_argument('--start', type=str, required=True,
                        help='开始日期（格式：YYYY-MM-DD）')
    parser.add_argument('--end', type=str, required=True,
                        help='结束日期（格式：YYYY-MM-DD）')
    parser.add_argument('--term', type=str, choices=['one_year', 'five_year'], default='one_year',
                        help='LPR期限（one_year: 一年期，five_year: 五年期以上）')
    parser.add_argument('--mag', type=int, default=1,
                        help='约定LPR倍数（默认：1倍）')
    parser.add_argument('--gap', type=str, choices=['no_tail', 'both'], default='no_tail',
                        help='天数算法：no_tail: 算头不算尾，both: 两头都算（默认：no_tail）')
    parser.add_argument('--day-count', type=int, choices=[360, 365], default=365,
                        help='计算年度天数：360或365（默认：365）')
    parser.add_argument('--update', action='store_true', help='强制更新LPR数据')
    parser.add_argument('--export', type=str,
                        help='导出Word文档的文件名（默认：LPR计算报告.docx）')
    parser.add_argument('--no-export', action='store_true', help='不导出Word文档')
    return parser.parse_args()


def main():
    # 解析命令行参数
    args = parse_arguments()

    # 如果指定更新或者文件不存在，则获取LPR数据
    import os
    if args.update or not os.path.exists('LPR_Data.txt'):
        print("正在获取LPR数据...")
        fetch_loan_rates()

    # 加载LPR数据
    lpr_data = load_lpr_data()
    if lpr_data is None:
        print("无法加载LPR数据，程序终止。")
        return

    # 准备计算参数
    amount = args.amount
    start_date = datetime.datetime.strptime(args.start, '%Y-%m-%d')
    end_date = datetime.datetime.strptime(args.end, '%Y-%m-%d')
    term = args.term
    mag = args.mag
    gap = args.gap
    day_count = args.day_count

    # 打印详细计算过程
    results = print_interest_details(
        amount, start_date, end_date, lpr_data, term, mag, gap, day_count)

    # 导出到Word文档
    if results and not args.no_export:
        try:
            # 获取当前时间并格式化为字符串
            current_time = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            
            # 构建输出文件名
            if args.export:
                # 如果用户提供了自定义文件名，在扩展名前添加时间戳
                filename, extension = os.path.splitext(args.export)
                output_file = f"{filename}_{current_time}{extension}"
            else:
                # 使用默认文件名加时间戳
                output_file = f"LPR计算报告_{current_time}.docx"
            
            export_to_word(results, output_file)
            print(f"报告已导出至: {output_file}")
        except Exception as e:
            print(f"导出Word文档时发生错误: {e}")
            print("请确保已安装python-docx库，可以通过 'pip install python-docx' 安装")


if __name__ == "__main__":
    main()